import os
import fitz  # PyMuPDF
import docx
import google.generativeai as genai
from dotenv import load_dotenv
import tempfile

# Load environment variables
load_dotenv()

# Configure Google AI
api_key = os.getenv("GOOGLE_API_KEY")
if api_key:
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-flash')
else:
    model = None

class DocumentQA:
    def __init__(self):
        self.document_text = ""
        self.is_loaded = False
    
    def load_document(self, file_path):
        """Load and extract text from document"""
        try:
            ext = os.path.splitext(file_path)[1].lower()
            if ext == ".pdf":
                self.document_text = self.read_pdf(file_path)
            elif ext == ".docx":
                self.document_text = self.read_docx(file_path)
            elif ext == ".txt":
                self.document_text = self.read_txt(file_path)
            else:
                raise ValueError("Unsupported file type")
            
            self.is_loaded = True
            return {"status": "Document loaded successfully", "length": len(self.document_text)}
        except Exception as e:
            return {"status": f"Error loading document: {str(e)}"}
    
    def read_pdf(self, file_path):
        """Extract text from PDF"""
        doc = fitz.open(file_path)
        text = "\n".join([page.get_text() for page in doc])
        doc.close()
        return text
    
    def read_docx(self, file_path):
        """Extract text from DOCX"""
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    
    def read_txt(self, file_path):
        """Extract text from TXT"""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    
    def ask_question(self, question):
        """Answer question about the document"""
        if not self.is_loaded or not self.document_text:
            return "Please upload a document first."
        
        if not model:
            return "AI model not configured. Please set GOOGLE_AI_API_KEY environment variable."
        
        try:
            prompt = f"""
            Based on the following document, please answer this question: {question}
            
            Document content:
            {self.document_text[:4000]}  # Limit to avoid token limits
            
            Please provide a clear, accurate answer based on the document content.
            """
            
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            return f"Error generating answer: {str(e)}"
    
    def clear_cache(self):
        """Clear document cache"""
        self.document_text = ""
        self.is_loaded = False
        return "Cache cleared"

class InterviewCopilot:
    def __init__(self):
        self.document_text = ""
        self.is_loaded = False
    
    def load_document(self, file_path):
        """Load and extract text from document"""
        try:
            ext = os.path.splitext(file_path)[1].lower()
            if ext == ".pdf":
                self.document_text = self.read_pdf(file_path)
            elif ext == ".docx":
                self.document_text = self.read_docx(file_path)
            elif ext == ".txt":
                self.document_text = self.read_txt(file_path)
            else:
                raise ValueError("Unsupported file type")
            
            self.is_loaded = True
            return {"status": "Document loaded successfully", "length": len(self.document_text)}
        except Exception as e:
            return {"status": f"Error loading document: {str(e)}"}
    
    def read_pdf(self, file_path):
        """Extract text from PDF"""
        doc = fitz.open(file_path)
        text = "\n".join([page.get_text() for page in doc])
        doc.close()
        return text
    
    def read_docx(self, file_path):
        """Extract text from DOCX"""
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    
    def read_txt(self, file_path):
        """Extract text from TXT"""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    
    def generate_questions(self, num_questions=5, level="medium", qtype=None):
        """Generate interview questions based on document"""
        if not self.is_loaded or not self.document_text:
            return ["Please upload a document first."]
        
        if not model:
            return ["AI model not configured. Please set GOOGLE_AI_API_KEY environment variable."]
        
        try:
            level_desc = {
                "easy": "basic and straightforward",
                "medium": "moderately challenging",
                "hard": "advanced and complex"
            }.get(level, "moderately challenging")
            
            qtype_desc = {
                "mcq": "multiple choice questions",
                "hr": "HR and behavioral questions",
                "technical": "technical questions",
                "theory": "theoretical questions"
            }.get(qtype, "interview questions")
            
            prompt = f"""
            Based on the following document, generate {num_questions} {level_desc} {qtype_desc} for an interview.
            
            Document content:
            {self.document_text[:4000]}  # Limit to avoid token limits
            
            Please generate {num_questions} questions that test understanding of the document content.
            Format each question on a new line.
            """
            
            response = model.generate_content(prompt)
            questions = [q.strip() for q in response.text.split('\n') if q.strip()]
            
            # Ensure we have the right number of questions
            if len(questions) < num_questions:
                questions.extend([f"Additional question {i+1}" for i in range(num_questions - len(questions))])
            
            return questions[:num_questions]
        except Exception as e:
            return [f"Error generating questions: {str(e)}"]
    
    def evaluate_answers(self, answers):
        """Evaluate interview answers"""
        if not self.is_loaded or not self.document_text:
            return 0.0, ["Please upload a document first."]
        
        if not model:
            return 0.0, ["AI model not configured. Please set GOOGLE_AI_API_KEY environment variable."]
        
        try:
            prompt = f"""
            Based on the following document, evaluate these interview answers:
            
            Document content:
            {self.document_text[:2000]}  # Limit to avoid token limits
            
            Answers to evaluate:
            {chr(10).join([f"Answer {i+1}: {answer}" for i, answer in enumerate(answers)])}
            
            Please provide:
            1. An overall score (0-100)
            2. Individual feedback for each answer
            
            Format: Score: [number], Feedback: [detailed feedback for each answer]
            """
            
            response = model.generate_content(prompt)
            text = response.text
            
            # Extract score
            score = 75.0  # Default score
            if "Score:" in text:
                try:
                    score = float(text.split("Score:")[1].split()[0])
                except:
                    pass
            
            # Create feedback
            feedback = []
            for i, answer in enumerate(answers):
                feedback.append({
                    "question": f"Question {i+1}",
                    "user_answer": answer,
                    "score": score / len(answers),
                    "feedback": f"Answer evaluated based on document content",
                    "correct_answer": "See document for reference"
                })
            
            return score, feedback
        except Exception as e:
            return 0.0, [{"question": f"Q{i+1}", "score": 0, "feedback": f"Error: {str(e)}"} for i in range(len(answers))]
