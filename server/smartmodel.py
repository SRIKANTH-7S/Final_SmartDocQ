# import os
# import fitz  # PyMuPDF
# import docx
# import faiss
# import numpy as np
# from sentence_transformers import SentenceTransformer
# import google.generativeai as genai
# from dotenv import load_dotenv

# # ================================
# # Step 1: Load environment variables and configure API
# # ================================
# load_dotenv()
# api_key = os.getenv("GOOGLE_API_KEY")
# if not api_key:
#     raise ValueError("GOOGLE_API_KEY not found in environment variables.")
# genai.configure(api_key=api_key)

# # ================================
# # Step 2: Text Extraction Functions
# # ================================
# def read_pdf(file_path):
#     doc = fitz.open(file_path)
#     return "\n".join([page.get_text() for page in doc])

# def read_docx(file_path):
#     doc = docx.Document(file_path)
#     return "\n".join([para.text for para in doc.paragraphs])

# def read_txt(file_path):
#     with open(file_path, "r", encoding="utf-8") as f:
#         return f.read()

# def read_document(file_path):
#     ext = os.path.splitext(file_path)[1].lower()
#     if ext == ".pdf":
#         return read_pdf(file_path)
#     elif ext == ".docx":
#         return read_docx(file_path)
#     elif ext == ".txt":
#         return read_txt(file_path)
#     else:
#         raise ValueError("Unsupported file type")

# # ================================
# # Step 3: Chunking Function
# # ================================
# def chunk_text(text, chunk_size=800, overlap=100):  # Increased size, overlap
#     words = text.split()
#     chunks = []
#     start = 0
#     while start < len(words):
#         end = min(start + chunk_size, len(words))
#         chunk = " ".join(words[start:end])
#         if chunk.strip():  # Only add non-empty chunks
#             chunks.append(chunk)
#         start += chunk_size - overlap
#     print(f"DEBUG: Chunked into {len(chunks)} valid chunks. Sizes: {[len(c.split()) for c in chunks]}")
#     return chunks
# def load_document(self, file_path):
#     text = read_document(file_path)
#     print(f"DEBUG: Raw text length: {len(text)} characters")
#     self.chunks = chunk_text(text)
#     valid_chunks = [c for c in self.chunks if c.strip()]
#     if len(valid_chunks) != len(self.chunks):
#         print(f"DEBUG: Warning: {len(self.chunks) - len(valid_chunks)} empty chunks filtered!")
#     self.chunks = valid_chunks
#     self.index, self.chunks = build_faiss_index(self.chunks)
#     return f"✅ Document loaded. Total chunks: {len(self.chunks)}"

# # ================================
# # Step 4: Embeddings & FAISS
# # ================================
# model = SentenceTransformer("all-MiniLM-L6-v2")

# def get_embeddings(chunks):
#     return model.encode(chunks, show_progress_bar=False)

# def build_faiss_index(chunks):
#     embeddings = get_embeddings(chunks)
#     print(f"DEBUG: Embedding shape: {embeddings.shape if hasattr(embeddings, 'shape') else 'None'}")
#     if not embeddings.size:
#         raise ValueError("DEBUG: No embeddings generated - check chunk content!")
#     embedding_array = np.array(embeddings).astype("float32")
#     index = faiss.IndexFlatL2(embedding_array.shape[1])
#     index.add(embedding_array)
#     print(f"DEBUG: FAISS index built with {index.ntotal} vectors")
#     return index, chunks

# def search_similar_chunks(query, index, chunks, model, k=3, threshold=1.2):
#     query_embedding = model.encode([query]).astype("float32")
#     distances, indices = index.search(query_embedding, k)

#     if distances[0][0] > threshold:
#         return None  
#     return [chunks[i] for i in indices[0]]


# ###########temporary####################
# # In read_document function (after ext check), add:
# def read_document(file_path):
#     ext = os.path.splitext(file_path)[1].lower()
#     if ext == ".pdf":
#         text = read_pdf(file_path)
#         print(f"DEBUG: Extracted PDF text preview: {text[:200]}...")  # First 200 chars
#         return text
#     # ... (similar for docx/txt, with print)

# # In load_document (DocumentQA class), after chunking:
# def load_document(self, file_path):
#     text = read_document(file_path)
#     self.chunks = chunk_text(text)
#     print(f"DEBUG: Created {len(self.chunks)} chunks. Sample: {self.chunks[0][:100]}...")  # Preview first chunk
#     self.index, self.chunks = build_faiss_index(self.chunks)
#     return f"✅ Document loaded. Total chunks: {len(self.chunks)}"

# # In search_similar_chunks, before return:
# def search_similar_chunks(query, index, chunks, model, k=10, threshold=1.5):
#     query_embedding = model.encode([query]).astype("float32")
#     distances, indices = index.search(query_embedding, k)
#     print(f"DEBUG: Query '{query}' - Top distances: {distances[0][:3]} (lower = better match)")  # Log top scores

#     if distances[0][0] > threshold:
#         print(f"DEBUG: No chunks above threshold {threshold} - Falling back to 'Sorry...'")
#         return None  
#     return [chunks[i] for i in indices[0]]

# # In ask_question, before Gemini:
# def ask_question(self, query):
#     if self.index is None or self.chunks is None:
#         return "⚠️ Please load a document first."
#     # Normalize query (simple fix for typos)
#     query = query.lower().replace("intresting", "interesting")
#     top_chunks = search_similar_chunks(query, self.index, self.chunks, model, k=3)
#     if top_chunks:
#         print(f"DEBUG: Retrieved {len(top_chunks)} chunks for query: {top_chunks[0][:100]}...")
#     else:
#         print("DEBUG: No chunks retrieved - using fallback.")
#     return ask_gemini_with_context(query, top_chunks)

# def load_document(self, file_path):
#     text = read_document(file_path)
#     self.chunks = chunk_text(text)
#     print(f"DEBUG: Created {len(self.chunks)} chunks. Sample: {self.chunks[0][:100]}...")  # Already there
#     # New: Check if chunks are valid
#     valid_chunks = [c for c in self.chunks if c.strip()]
#     if len(valid_chunks) != len(self.chunks):
#         print(f"DEBUG: Warning: {len(self.chunks) - len(valid_chunks)} empty chunks filtered out!")
#     self.chunks = valid_chunks
#     self.index, self.chunks = build_faiss_index(self.chunks)
#     return f"✅ Document loaded. Total chunks: {len(self.chunks)}"

# def build_faiss_index(chunks):
#     embeddings = get_embeddings(chunks)
#     print(f"DEBUG: Embedding shape: {embeddings.shape if hasattr(embeddings, 'shape') else 'None'}")  # Check dims
#     if not embeddings.size:  # Empty array
#         raise ValueError("DEBUG: No embeddings generated - check chunk content!")
#     embedding_array = np.array(embeddings).astype("float32")
#     index = faiss.IndexFlatL2(embedding_array.shape[1])
#     index.add(embedding_array)
#     print(f"DEBUG: FAISS index built with {index.ntotal} vectors")
#     return index, chunks

# ##########temp

# # ================================
# # Step 5: Ask Gemini (Strict Context Only)
# # ================================
# def ask_gemini_with_context(query, context_chunks=None):
#     try:
#         model_name = "gemini-1.5-flash"  
#         gmodel = genai.GenerativeModel(model_name=model_name)

#         if not context_chunks:
#             return "Sorry, I cannot answer this question as it is not in the document."

#         context_text = "\n".join(context_chunks)
#         prompt = f"""
#         You are a document Q&A assistant. Use the context below to answer. 
#         If the answer is clear, respond directly. If partial, say: 'Based on the document, [summarize]. For details, ask about a specific section.'
        
#         Context:
#         {context_text}

#         Question:
#         {query}

#         Answer:
#         """
#         response = gmodel.generate_content(prompt)
#         return response.text.strip()

#     except Exception as e:
#         return f"An error occurred in Gemini API: {str(e)}"
# # ================================
# # Step 6: API-like usage (No Menu)
# # ================================
# class DocumentQA:
#     def __init__(self):
#         self.index = None
#         self.chunks = None

#     def load_document(self, file_path):
#         text = read_document(file_path)
#         self.chunks = chunk_text(text)
#         self.index, self.chunks = build_faiss_index(self.chunks)
#         return f"✅ Document loaded. Total chunks: {len(self.chunks)}"

#     def ask_question(self, query):
#         if self.index is None or self.chunks is None:
#             return "⚠️ Please load a document first."
#         top_chunks = search_similar_chunks(query, self.index, self.chunks, model, k=5)
#         return ask_gemini_with_context(query, top_chunks)

#     def clear_cache(self):
#         """Clear RAM cache (for logout)"""
#         self.index = None
#         self.chunks = None
#         return "Cache cleared."




# smartmodel.py
import os
from typing import List, Optional, Tuple
import fitz  # PyMuPDF
import docx
import numpy as np
import faiss
from sentence_transformers import SentenceTransformer
import google.generativeai as genai
from dotenv import load_dotenv
import json
import os

load_dotenv()
API_KEY = os.getenv("GOOGLE_API_KEY")
if not API_KEY:
    raise ValueError("GOOGLE_API_KEY not set in environment")
genai.configure(api_key=API_KEY)

# Global embedding model (one-time load)
EMBED_MODEL_NAME = "all-MiniLM-L6-v2"
_embed_model = SentenceTransformer(EMBED_MODEL_NAME)


# ---------------------
# Document reading
# ---------------------
def read_pdf(path: str) -> str:
    doc = fitz.open(path)
    pages = [p.get_text() for p in doc]
    return "\n".join(pages)


def read_docx(path: str) -> str:
    d = docx.Document(path)
    return "\n".join([p.text for p in d.paragraphs])


def read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def read_document_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return read_pdf(path)
    if ext == ".docx":
        return read_docx(path)
    if ext == ".txt":
        return read_txt(path)
    raise ValueError("Unsupported file type: " + ext)


# ---------------------
# Chunking
# ---------------------
def chunk_text(text: str, chunk_size: int = 300, overlap: int = 50) -> List[str]:
    """Chunk by words. chunk_size is number of words per chunk."""
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        j = min(i + chunk_size, len(words))
        chunk = " ".join(words[i:j]).strip()
        if chunk:
            chunks.append(chunk)
        i += chunk_size - overlap
    return chunks


# ---------------------
# Embeddings & FAISS
# ---------------------
def get_embeddings(texts: List[str]) -> np.ndarray:
    arr = _embed_model.encode(texts, show_progress_bar=False)
    # SentenceTransformer may return numpy array already
    if isinstance(arr, list):
        arr = np.array(arr)
    return np.asarray(arr).astype("float32")


def build_faiss_index(chunks: List[str]) -> Tuple[faiss.IndexFlatL2, List[str]]:
    if not chunks:
        raise ValueError("No chunks to index")
    embeddings = get_embeddings(chunks)
    if embeddings.size == 0:
        raise ValueError("Embeddings returned empty")
    dim = embeddings.shape[1]
    index = faiss.IndexFlatL2(dim)
    index.add(embeddings)
    return index, chunks


def search_similar_chunks(
    query: str,
    index: faiss.IndexFlatL2,
    chunks: List[str],
    top_k: int = 5,
) -> List[str]:
    if index is None or index.ntotal == 0:
        return []
    q_emb = get_embeddings([query])  # shape (1, dim)
    distances, indices = index.search(q_emb, top_k)
    # distances are L2 distances (lower = better). We will return top_k regardless,
    # but filter out invalid indices (-1)
    result = []
    for idx in indices[0]:
        if idx == -1:
            continue
        if 0 <= idx < len(chunks):
            result.append(chunks[idx])
    return result


def ask_gemini_with_context(query: str, context_chunks: List[str]) -> str:
    """Ask Gemini a question using only the provided context chunks. Implements jittered exponential backoff on quota errors."""
    if not context_chunks:
        return "Sorry — I cannot find the answer in the document."

    context_text = "\n\n---\n\n".join(context_chunks)
    prompt = f"""
You are a document Q&A assistant. Answer concisely and only using facts from the provided context. If the answer is not in the context, say: "I cannot find the answer in the document.".

Context:
{context_text}

Question:
{query}

Answer:
"""

    import time
    import random
    import logging
    import re

    logger = logging.getLogger("gemini_retry")
    if not logger.handlers:
        fh = logging.FileHandler(os.path.join(os.path.dirname(__file__), "gemini_retry.log"))
        fh.setFormatter(logging.Formatter("%(asctime)s %(levelname)s %(message)s"))
        logger.addHandler(fh)
        logger.setLevel(logging.INFO)

    max_retries = 5
    backoff = 1.0
    for attempt in range(max_retries):
        try:
            model = genai.GenerativeModel(model_name="gemini-2.0-flash")
            response = model.generate_content(prompt)
            text = ""
            if hasattr(response, "text") and response.text:
                text = response.text
            elif hasattr(response, "candidates") and response.candidates:
                text = response.candidates[0].content
            else:
                text = str(response)
            # best-effort cache of successful response
            try:
                import hashlib, threading
                cache_dir = os.path.join(os.path.dirname(__file__), ".cache")
                os.makedirs(cache_dir, exist_ok=True)
                cache_path = os.path.join(cache_dir, "gemini_cache.json")
                prompt_key = hashlib.sha256(prompt.encode("utf-8")).hexdigest()
                with threading.Lock():
                    try:
                        with open(cache_path, "r", encoding="utf-8") as cf:
                            cache = json.load(cf)
                    except Exception:
                        cache = {}
                    cache[prompt_key] = text.strip()
                    try:
                        with open(cache_path, "w", encoding="utf-8") as cf:
                            json.dump(cache, cf, ensure_ascii=False, indent=2)
                    except Exception:
                        pass
            except Exception:
                pass
            return text.strip()
        except Exception as e:
            err_str = repr(e)
            if "ResourceExhausted" in err_str or "quota" in err_str.lower():
                retry_seconds = None
                try:
                    m = re.search(r"retry in\s*([0-9]+\.?[0-9]*)s", err_str, re.IGNORECASE)
                    if m:
                        retry_seconds = float(m.group(1))
                except Exception:
                    retry_seconds = None

                if attempt < max_retries - 1:
                    sleep_time = retry_seconds if retry_seconds is not None else backoff * (0.5 + random.random())
                    sleep_time = min(sleep_time, 60)
                    logger.info(f"Gemini quota error detected, attempt {attempt+1}/{max_retries}, sleeping {sleep_time}s before retry. Error: {err_str}")
                    time.sleep(sleep_time)
                    backoff *= 2
                    continue
                else:
                    logger.warning(f"Gemini quota exhausted after {max_retries} attempts. Last error: {err_str}")
                    # fallback to cache
                    try:
                        import hashlib, threading
                        cache_dir = os.path.join(os.path.dirname(__file__), ".cache")
                        cache_path = os.path.join(cache_dir, "gemini_cache.json")
                        prompt_key = hashlib.sha256(prompt.encode("utf-8")).hexdigest()
                        with threading.Lock():
                            with open(cache_path, "r", encoding="utf-8") as cf:
                                cache = json.load(cf)
                            if prompt_key in cache:
                                logger.info("Returning cached Gemini response due to quota exhaustion")
                                return f"(cached) {cache[prompt_key]}"
                    except Exception:
                        pass
                    return "Service temporarily unavailable due to API quota limits. Please try again later."
            logger.error(f"Gemini call failed: {err_str}")
            return f"Gemini error: {err_str}"


# ---------------------
# Public DocumentQA class
# ---------------------
class DocumentQA:
    def __init__(self):
        self.index: Optional[faiss.IndexFlatL2] = None
        self.chunks: Optional[List[str]] = None

    def load_document(self, file_path: str) -> str:
        text = read_document_file(file_path)
        if not text or not text.strip():
            raise ValueError("Document empty after extraction")
        chunks = chunk_text(text, chunk_size=300, overlap=50)
        if not chunks:
            raise ValueError("No chunks created from document")
        self.index, self.chunks = build_faiss_index(chunks)
        return f"Document loaded. {len(self.chunks)} chunks indexed."

    def save_index(self, dest_dir: str) -> None:
        """Save the FAISS index and chunks to dest_dir. Creates dest_dir if missing."""
        if not self.index or not self.chunks:
            raise ValueError("No index/chunks to save")
        os.makedirs(dest_dir, exist_ok=True)
        index_path = os.path.join(dest_dir, "index.faiss")
        chunks_path = os.path.join(dest_dir, "chunks.json")
        # write faiss index
        faiss.write_index(self.index, index_path)
        # write chunks
        with open(chunks_path, "w", encoding="utf-8") as f:
            json.dump(self.chunks, f, ensure_ascii=False)

    @staticmethod
    def load_index_from(dir_path: str) -> 'DocumentQA':
        """Load a persisted FAISS index and chunks from dir_path and return a DocumentQA instance."""
        index_path = os.path.join(dir_path, "index.faiss")
        chunks_path = os.path.join(dir_path, "chunks.json")
        if not os.path.exists(index_path) or not os.path.exists(chunks_path):
            raise FileNotFoundError("Index or chunks not found in provided directory")
        # read chunks
        with open(chunks_path, "r", encoding="utf-8") as f:
            chunks = json.load(f)
        # load index
        index = faiss.read_index(index_path)
        inst = DocumentQA()
        inst.index = index
        inst.chunks = chunks
        return inst

    def ask_question(self, query: str) -> str:
        if not query or not query.strip():
            return "Please ask a valid question."
        if self.index is None or self.chunks is None:
            return "⚠️ Please load a document first."
        top_chunks = search_similar_chunks(query, self.index, self.chunks, top_k=5)
        if not top_chunks:
            return "Sorry — I cannot find the answer in the document."
        # Send to Gemini with strict context
        return ask_gemini_with_context(query, top_chunks)

    def clear_cache(self) -> str:
        self.index = None
        self.chunks = None
        return "Cache cleared"
