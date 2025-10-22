import os
import fitz  # PyMuPDF
import docx
import google.generativeai as genai
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure Google AI - Simple and reliable
api_key = os.getenv("GOOGLE_AI_API_KEY")
model = None

if api_key:
    try:
        genai.configure(api_key=api_key)
        # Use the most basic model that should work
        model = genai.GenerativeModel('gemini-pro')
        print(f"✅ Google AI configured successfully")
    except Exception as e:
        print(f"❌ Error configuring Google AI: {e}")
        model = None
else:
    print("❌ GOOGLE_AI_API_KEY not found")

class DocumentQA:
    def __init__(self):
        self.document_text = ""
        self.is_loaded = False
    
    def load_document(self, file_path):
        """Load and extract text from document"""
        try:
            ext = os.path.splitext(file_path)[1].lower()
            
            if ext == ".pdf":
                self.document_text = self.read_pdf(file_path)
            elif ext == ".docx":
                self.document_text = self.read_docx(file_path)
            elif ext == ".txt":
                self.document_text = self.read_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {ext}")
            
            self.is_loaded = True
            return {"status": "Document loaded successfully", "length": len(self.document_text)}
        except Exception as e:
            return {"status": f"Error loading document: {str(e)}"}
    
    def read_pdf(self, file_path):
        """Extract text from PDF"""
        doc = fitz.open(file_path)
        text = "\n".join([page.get_text() for page in doc])
        doc.close()
        return text
    
    def read_docx(self, file_path):
        """Extract text from DOCX"""
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    
    def read_txt(self, file_path):
        """Extract text from TXT"""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    
    def ask_question(self, question):
        """Answer question about the document"""
        if not self.is_loaded or not self.document_text:
            return "Please upload a document first."
        
        if not model:
            return "AI model not configured. Please set GOOGLE_AI_API_KEY environment variable."
        
        try:
            # Simple prompt with limited text
            doc_text = self.document_text[:2000]  # Limit text
            
            prompt = f"""Based on this document, answer: {question}

Document: {doc_text}

Answer:"""
            
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            return f"Error: {str(e)}"
    
    def clear_cache(self):
        """Clear document cache"""
        self.document_text = ""
        self.is_loaded = False
        return "Cache cleared"

class InterviewCopilot:
    def __init__(self):
        self.document_text = ""
        self.is_loaded = False
    
    def load_document(self, file_path):
        """Load and extract text from document"""
        try:
            ext = os.path.splitext(file_path)[1].lower()
            
            if ext == ".pdf":
                self.document_text = self.read_pdf(file_path)
            elif ext == ".docx":
                self.document_text = self.read_docx(file_path)
            elif ext == ".txt":
                self.document_text = self.read_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {ext}")
            
            self.is_loaded = True
            return {"status": "Document loaded successfully", "length": len(self.document_text)}
        except Exception as e:
            return {"status": f"Error loading document: {str(e)}"}
    
    def read_pdf(self, file_path):
        """Extract text from PDF"""
        doc = fitz.open(file_path)
        text = "\n".join([page.get_text() for page in doc])
        doc.close()
        return text
    
    def read_docx(self, file_path):
        """Extract text from DOCX"""
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    
    def read_txt(self, file_path):
        """Extract text from TXT"""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    
    def generate_questions(self, num_questions=5, level="medium", qtype=None):
        """Generate interview questions based on document"""
        if not self.is_loaded or not self.document_text:
            return ["Please upload a document first."]
        
        if not model:
            return ["AI model not configured. Please set GOOGLE_AI_API_KEY environment variable."]
        
        try:
            doc_text = self.document_text[:1500]  # Limit text
            
            prompt = f"""Based on this document, generate {num_questions} interview questions:

Document: {doc_text}

Questions:"""
            
            response = model.generate_content(prompt)
            questions_text = response.text
            questions = [q.strip() for q in questions_text.split('\n') if q.strip()]
            
            # Ensure we have enough questions
            if len(questions) < num_questions:
                questions.extend([f"Question {i+1}" for i in range(num_questions - len(questions))])
            
            return questions[:num_questions]
        except Exception as e:
            return [f"Error: {str(e)}"]
    
    def evaluate_answers(self, answers):
        """Evaluate interview answers"""
        if not self.is_loaded or not self.document_text:
            return 0.0, ["Please upload a document first."]
        
        if not model:
            return 0.0, ["AI model not configured. Please set GOOGLE_AI_API_KEY environment variable."]
        
        try:
            doc_text = self.document_text[:1000]  # Limit text
            
            prompt = f"""Evaluate these answers based on the document:

Document: {doc_text}

Answers: {', '.join(answers)}

Score (0-100):"""
            
            response = model.generate_content(prompt)
            text = response.text
            
            # Extract score
            score = 75.0  # Default
            if "Score:" in text:
                try:
                    score = float(text.split("Score:")[1].split()[0])
                except:
                    pass
            
            # Create feedback
            feedback = []
            for i, answer in enumerate(answers):
                feedback.append({
                    "question": f"Question {i+1}",
                    "user_answer": answer,
                    "score": score / len(answers),
                    "feedback": "Evaluated by AI",
                    "correct_answer": "See document"
                })
            
            return score, feedback
        except Exception as e:
            return 0.0, [{"question": f"Q{i+1}", "score": 0, "feedback": f"Error: {str(e)}"} for i in range(len(answers))]
